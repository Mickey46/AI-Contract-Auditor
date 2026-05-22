"""Parse DOCX files into section-aware chunks with metadata."""

from __future__ import annotations
from docx import Document
from app.models.schemas import Chunk


def parse_docx(file_path: str) -> list[Chunk]:
    chunks: list[Chunk] = []
    filename = file_path.split("/")[-1]
    doc = Document(file_path)

    current_section = "Introduction"
    current_lines: list[str] = []
    chunk_index = 0

    def flush(section: str, lines: list[str], idx: int) -> Chunk | None:
        text = "\n".join(lines).strip()
        if not text:
            return None
        return Chunk(
            text=f"[Section: {section}]\n{text}",
            source_file=filename,
            source_type="docx",
            section=section,
            chunk_index=idx,
            doc_precedence=2,
        )

    for para in doc.paragraphs:
        style = para.style.name or ""
        text = para.text.strip()

        if style.startswith("Heading") and text:
            # Flush previous section
            c = flush(current_section, current_lines, chunk_index)
            if c:
                chunks.append(c)
                chunk_index += 1
            current_section = text
            current_lines = []
        elif text:
            current_lines.append(text)

    # Flush last section
    c = flush(current_section, current_lines, chunk_index)
    if c:
        chunks.append(c)

    # Also extract table data
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text = " | ".join(cells)
            if any(cells):
                chunks.append(
                    Chunk(
                        text=f"[Table in {filename}]: {text}",
                        source_file=filename,
                        source_type="docx",
                        section="Table",
                        chunk_index=chunk_index,
                        doc_precedence=2,
                    )
                )
                chunk_index += 1

    return chunks
