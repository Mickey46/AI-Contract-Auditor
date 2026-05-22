"""
Generates a complete realistic synthetic dataset for the AI Contract Auditor:
  - master_contract.pdf       (base contract, PDF, pages with pricing tables)
  - pricing.xlsx              (3 sheets: Base Pricing, Volume Tiers, Discounts)
  - amendment_q2.docx         (overrides CP-001 price; changes EL-003 discount)
  - email_addendum.eml        (adds extra discount for DM-004)
  - invoice_INV-1001.csv      (intentional mismatches on CP-001 and EL-003)
"""

import os
import sys

# Allow running from repo root
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)
from reportlab.lib import colors
from reportlab.lib.units import inch
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor
import csv
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "contracts")
INV_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "invoices")
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(INV_DIR, exist_ok=True)


# ── Master contract terms (ground truth) ──────────────────────────────────────
BASE_TERMS = {
    "CP-001": {"desc": "Claims Processing",   "unit_price": 5.50, "discount": 5,  "tax": 8},
    "AP-002": {"desc": "Appeals Processing",  "unit_price": 12.00, "discount": 0,  "tax": 8},
    "EL-003": {"desc": "Eligibility Checks",  "unit_price": 3.25, "discount": 8,  "tax": 8},
    "DM-004": {"desc": "Data Migration",      "unit_price": 45.00, "discount": 0,  "tax": 8},
}

# Amendment overrides (amendment_q2.docx)
AMENDMENT_TERMS = {
    "CP-001": {"unit_price": 5.00, "discount": 10},   # price drops, discount rises
    "EL-003": {"discount": 12},                         # only discount changes
}

# Email addendum override
EMAIL_ADDENDUM = {
    "DM-004": {"discount": 5},                          # new 5% discount added
}

# Final authoritative values (what the invoice SHOULD match)
FINAL_TERMS = {sku: dict(BASE_TERMS[sku]) for sku in BASE_TERMS}
for sku, overrides in AMENDMENT_TERMS.items():
    FINAL_TERMS[sku].update(overrides)
for sku, overrides in EMAIL_ADDENDUM.items():
    FINAL_TERMS[sku].update(overrides)


# ── 1. PDF ─────────────────────────────────────────────────────────────────────
def make_pdf():
    path = os.path.join(OUT_DIR, "master_contract.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter,
                            rightMargin=inch * 0.75, leftMargin=inch * 0.75,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]

    story = []

    def para(text, style=body):
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 6))

    para("MASTER SERVICES AGREEMENT", h1)
    para("Agreement No.: MSA-2025-CVS-0047", body)
    para("Effective Date: January 1, 2025", body)
    para('Parties: CVS Health Innovation Lab ("Client") and Apex Health Solutions ("Provider")', body)
    story.append(Spacer(1, 12))

    para("1. OVERVIEW", h2)
    para(
        "This Master Services Agreement (\"Agreement\") governs the provision of healthcare "
        "technology services by Apex Health Solutions to CVS Health Innovation Lab. All pricing, "
        "discount schedules, and billing rules are defined herein unless superseded by a written amendment.",
        body,
    )

    para("2. SERVICES AND BASE PRICING", h2)
    para(
        "The following table defines the base unit pricing for each Service SKU. "
        "Prices are in USD per transaction unless otherwise stated.",
        body,
    )

    table_data = [
        ["SKU", "Service Description", "Unit Price (USD)", "Standard Discount (%)", "Tax Rate (%)"],
    ]
    for sku, t in BASE_TERMS.items():
        table_data.append([sku, t["desc"], f"${t['unit_price']:.2f}", f"{t['discount']}%", f"{t['tax']}%"])

    tbl = Table(table_data, colWidths=[0.9 * inch, 2.2 * inch, 1.3 * inch, 1.5 * inch, 1.0 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2FF")]),
        ("ALIGN", (2, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))

    para("3. DISCOUNT POLICY", h2)
    para(
        "Standard discounts are applied as defined in Section 2. Volume discounts are additive "
        "and apply when monthly transaction counts exceed thresholds defined in the Pricing Schedule "
        "(Exhibit A — pricing.xlsx). Discounts are applied to the pre-tax unit price.",
        body,
    )

    para("4. TAX POLICY", h2)
    para(
        "A uniform tax rate of 8% applies to all services unless the Client provides a valid "
        "tax-exemption certificate. Tax is computed on the discounted unit price multiplied by quantity.",
        body,
    )

    para("5. TOTAL AMOUNT CALCULATION", h2)
    para(
        "Total Amount = Quantity × Unit Price × (1 − Discount%) × (1 + Tax%)",
        body,
    )

    para("6. INVOICING", h2)
    para(
        "Invoices must be submitted monthly using the standard CSV format. Each line item must "
        "reference the SKU, quantity, unit price, applied discount, tax rate, and computed total. "
        "Invoices referencing incorrect unit prices or unapproved discounts will be rejected.",
        body,
    )

    para("7. AMENDMENTS", h2)
    para(
        "Amendments to this Agreement must be in writing and signed by both parties. "
        "An executed amendment supersedes conflicting terms in this Agreement or any prior exhibit. "
        "Amendments take effect on the date specified therein.",
        body,
    )

    story.append(PageBreak())

    para("EXHIBIT A — PRICING SCHEDULE REFERENCE", h1)
    para(
        "The full Pricing Schedule, including tiered volume discount thresholds, is maintained "
        "in pricing.xlsx (attached). The Excel file constitutes Exhibit A and is incorporated "
        "by reference into this Agreement.",
        body,
    )

    para("SIGNATURE PAGE", h2)
    para("CVS Health Innovation Lab: ________________________    Date: ___________", body)
    para("Apex Health Solutions:     ________________________    Date: ___________", body)

    doc.build(story)
    print(f"  Created: {path}")


# ── 2. Excel ───────────────────────────────────────────────────────────────────
def make_excel():
    path = os.path.join(OUT_DIR, "pricing.xlsx")
    wb = openpyxl.Workbook()

    header_fill = PatternFill("solid", fgColor="003366")
    subheader_fill = PatternFill("solid", fgColor="5B8DB8")
    alt_fill = PatternFill("solid", fgColor="EEF2FF")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    subheader_font = Font(bold=True, color="FFFFFF", size=9)
    bold = Font(bold=True)
    thin = Side(style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Sheet 1: Base Pricing ──────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Base Pricing"

    ws1.merge_cells("A1:G1")
    ws1["A1"] = "APEX HEALTH SOLUTIONS — BASE PRICING SCHEDULE"
    ws1["A1"].font = Font(bold=True, size=13, color="003366")
    ws1["A1"].alignment = Alignment(horizontal="center")

    ws1.merge_cells("A2:G2")
    ws1["A2"] = "Effective: 2025-01-01 | Agreement: MSA-2025-CVS-0047 | Currency: USD"
    ws1["A2"].alignment = Alignment(horizontal="center")
    ws1["A2"].font = Font(size=9, italic=True, color="555555")

    headers = ["SKU", "Service Description", "Category", "Unit Price", "Std Discount (%)", "Tax Rate (%)", "Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    rows = [
        ("CP-001", "Claims Processing",  "Transaction",  5.50,  5, 8, "Standard per-claim fee"),
        ("AP-002", "Appeals Processing", "Transaction", 12.00,  0, 8, "No discount — premium SLA"),
        ("EL-003", "Eligibility Checks", "Lookup",       3.25,  8, 8, "Bulk lookup service"),
        ("DM-004", "Data Migration",     "Project",     45.00,  0, 8, "Per-record migration fee"),
    ]
    for r, row in enumerate(rows, 4):
        fill = alt_fill if r % 2 == 0 else None
        for c, val in enumerate(row, 1):
            cell = ws1.cell(row=r, column=c, value=val)
            cell.border = border
            if fill:
                cell.fill = fill
            if c == 4:
                cell.number_format = '"$"#,##0.00'
            if c in (5, 6):
                cell.number_format = '0"%"'

    ws1.column_dimensions["A"].width = 10
    ws1.column_dimensions["B"].width = 26
    ws1.column_dimensions["C"].width = 14
    ws1.column_dimensions["D"].width = 13
    ws1.column_dimensions["E"].width = 17
    ws1.column_dimensions["F"].width = 13
    ws1.column_dimensions["G"].width = 30

    # ── Sheet 2: Volume Tiers ──────────────────────────────────────────────────
    ws2 = wb.create_sheet("Volume Tiers")

    ws2.merge_cells("A1:F1")
    ws2["A1"] = "VOLUME DISCOUNT TIERS — ADDITIVE ON TOP OF STANDARD DISCOUNT"
    ws2["A1"].font = Font(bold=True, size=12, color="003366")
    ws2["A1"].alignment = Alignment(horizontal="center")

    ws2.merge_cells("A2:F2")
    ws2["A2"] = "Volume thresholds are measured per SKU per calendar month. Tiers are mutually exclusive — only the highest applicable tier applies."
    ws2["A2"].font = Font(size=8, italic=True)
    ws2["A2"].alignment = Alignment(wrap_text=True)
    ws2.row_dimensions[2].height = 30

    # Multi-row header
    ws2.merge_cells("A3:A4"); ws2["A3"] = "SKU"
    ws2.merge_cells("B3:B4"); ws2["B3"] = "Service"
    ws2.merge_cells("C3:F3"); ws2["C3"] = "Monthly Volume Threshold — Additional Discount (%)"

    for c in range(1, 7):
        ws2.cell(row=3, column=c).fill = header_fill
        ws2.cell(row=3, column=c).font = header_font
        ws2.cell(row=3, column=c).alignment = Alignment(horizontal="center", vertical="center")
        ws2.cell(row=3, column=c).border = border
        ws2.cell(row=4, column=c).fill = subheader_fill
        ws2.cell(row=4, column=c).font = subheader_font
        ws2.cell(row=4, column=c).alignment = Alignment(horizontal="center")
        ws2.cell(row=4, column=c).border = border

    ws2["C4"] = "< 1,000 units"
    ws2["D4"] = "1,000 – 9,999 units"
    ws2["E4"] = "10,000 – 49,999 units"
    ws2["F4"] = "≥ 50,000 units"

    tier_rows = [
        ("CP-001", "Claims Processing",  0, 0, 2, 4),
        ("AP-002", "Appeals Processing", 0, 0, 1, 2),
        ("EL-003", "Eligibility Checks", 0, 1, 3, 5),
        ("DM-004", "Data Migration",     0, 0, 0, 2),
    ]
    for r, row in enumerate(tier_rows, 5):
        fill = alt_fill if r % 2 == 0 else None
        for c, val in enumerate(row, 1):
            cell = ws2.cell(row=r, column=c, value=val)
            cell.border = border
            if fill:
                cell.fill = fill

    for col in ["A", "B", "C", "D", "E", "F"]:
        ws2.column_dimensions[col].width = 22

    # ── Sheet 3: Discounts & Amendments Log ────────────────────────────────────
    ws3 = wb.create_sheet("Discounts & Amendments")

    ws3.merge_cells("A1:G1")
    ws3["A1"] = "DISCOUNT SCHEDULE & AMENDMENT LOG"
    ws3["A1"].font = Font(bold=True, size=12, color="003366")
    ws3["A1"].alignment = Alignment(horizontal="center")

    ws3.merge_cells("A2:G2")
    ws3["A2"] = "This sheet records all active discounts and amendment history. Amendments supersede base contract terms."
    ws3["A2"].font = Font(size=8, italic=True)

    disc_headers = ["SKU", "Service", "Discount Type", "Discount (%)", "Effective Date", "Expiry Date", "Source Document"]
    for col, h in enumerate(disc_headers, 1):
        cell = ws3.cell(row=3, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    disc_rows = [
        ("CP-001", "Claims Processing",  "Standard",           5,  "2025-01-01", "Ongoing",    "master_contract.pdf"),
        ("CP-001", "Claims Processing",  "Amendment Override", 10,  "2026-01-01", "Ongoing",    "amendment_q2.docx"),
        ("AP-002", "Appeals Processing", "Standard",           0,  "2025-01-01", "Ongoing",    "master_contract.pdf"),
        ("EL-003", "Eligibility Checks", "Standard",           8,  "2025-01-01", "2025-12-31", "master_contract.pdf"),
        ("EL-003", "Eligibility Checks", "Amendment Override", 12, "2026-01-01", "Ongoing",    "amendment_q2.docx"),
        ("DM-004", "Data Migration",     "Standard",           0,  "2025-01-01", "2025-12-31", "master_contract.pdf"),
        ("DM-004", "Data Migration",     "Email Addendum",     5,  "2026-01-01", "Ongoing",    "email_addendum.eml"),
    ]
    for r, row in enumerate(disc_rows, 4):
        fill = alt_fill if r % 2 == 0 else None
        for c, val in enumerate(row, 1):
            cell = ws3.cell(row=r, column=c, value=val)
            cell.border = border
            if fill:
                cell.fill = fill
            if c == 4:
                cell.number_format = '0"%"'

    for col in ["A", "B", "C", "D", "E", "F", "G"]:
        ws3.column_dimensions[col].width = 22

    wb.save(path)
    print(f"  Created: {path}")


# ── 3. DOCX Amendment ──────────────────────────────────────────────────────────
def make_docx():
    path = os.path.join(OUT_DIR, "amendment_q2.docx")
    doc = Document()

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0, 51, 102)

    def para(text):
        doc.add_paragraph(text)

    heading("AMENDMENT NO. 1 TO MASTER SERVICES AGREEMENT")
    para("Agreement Reference: MSA-2025-CVS-0047")
    para("Amendment Effective Date: January 1, 2026")
    para(
        'This Amendment No. 1 ("Amendment") is entered into by CVS Health Innovation Lab ("Client") '
        'and Apex Health Solutions ("Provider") and amends the Master Services Agreement dated '
        "January 1, 2025. Except as modified herein, all other terms of the Agreement remain in effect."
    )
    doc.add_paragraph()

    heading("Section 1 — Purpose", level=2)
    para(
        "Following renegotiation of service volumes and updated operational cost structures, "
        "the parties agree to revise the unit pricing and discount schedule for select SKUs "
        "effective Q1 2026."
    )

    heading("Section 2 — Revised Unit Pricing", level=2)
    para(
        "Effective January 1, 2026, the unit price for SKU CP-001 (Claims Processing) is "
        "REVISED from $5.50 to $5.00 per transaction. This amendment supersedes the base "
        "contract price stated in Section 2 of the Master Services Agreement and in "
        "pricing.xlsx (Base Pricing sheet)."
    )
    para(
        "All other SKU unit prices remain unchanged per the Master Services Agreement."
    )

    heading("Section 3 — Revised Discount Schedule", level=2)
    para(
        "Effective January 1, 2026, the following discount changes apply:"
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "SKU"
    hdr[1].text = "Service"
    hdr[2].text = "Previous Discount (%)"
    hdr[3].text = "New Discount (%)"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ("CP-001", "Claims Processing", "5%", "10%"),
        ("EL-003", "Eligibility Checks", "8%", "12%"),
    ]
    for row_data in data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()
    para(
        "The revised discounts are additive with any applicable volume tier discounts. "
        "Invoices submitted after January 1, 2026 must reflect the updated discount percentages."
    )

    heading("Section 4 — No Other Changes", level=2)
    para(
        "Except as expressly modified above, all terms and conditions of the Master Services "
        "Agreement MSA-2025-CVS-0047, including Exhibit A (pricing.xlsx), remain in full force."
    )

    heading("Section 5 — Signatures", level=2)
    para("CVS Health Innovation Lab: ________________________    Date: 2026-01-01")
    para("Apex Health Solutions:     ________________________    Date: 2026-01-01")

    doc.save(path)
    print(f"  Created: {path}")


# ── 4. Email Addendum ──────────────────────────────────────────────────────────
def make_email():
    path = os.path.join(OUT_DIR, "email_addendum.eml")

    body = """\
From: Diana Ross <diana.ross@apexhealthsolutions.com>
To: Prajwal Npraju <prajwal.n@cvs.com>; SarathChandra.Chennamsetty@aetna.com
CC: contracts@apexhealthsolutions.com
Subject: RE: MSA-2025-CVS-0047 — DM-004 Data Migration Discount Addendum (Q1 2026)
Date: Mon, 02 Dec 2025 10:14:37 -0500

Hi Prajwal,

As discussed in our call last Friday, we have agreed to extend a 5% discount on SKU DM-004
(Data Migration) effective January 1, 2026.

Summary of the agreed terms:

  SKU:              DM-004 — Data Migration
  Discount Change:  0% (current) → 5% (effective 2026-01-01)
  Unit Price:       $45.00 (unchanged)
  Tax Rate:         8% (unchanged)

This email serves as a binding addendum to Master Services Agreement MSA-2025-CVS-0047
and supersedes any prior discount schedule for DM-004. Please reference this email as
"Email Addendum — DM-004 Discount, dated 2025-12-02" in all future invoices and audits.

Please reply to confirm acceptance of these terms.

Best regards,
Diana Ross
VP of Contracts — Apex Health Solutions
Ph: +1 (404) 789-0023

---
This email and any attachments are confidential. If received in error, please notify the sender.
"""

    with open(path, "w") as f:
        f.write(body)

    print(f"  Created: {path}")


# ── 5. Invoice CSV ─────────────────────────────────────────────────────────────
def make_invoice():
    path = os.path.join(INV_DIR, "invoice_INV-1001.csv")

    # Intentional mismatches:
    # CP-001: uses old unit_price $5.50 (should be $5.00) and discount 5% (should be 10%)
    # EL-003: uses discount 8% (should be 12%)
    # DM-004: uses discount 0% (should be 5%)
    # AP-002: correct
    rows = [
        # invoice_id, date, sku, description, quantity, unit_price, discount_pct, tax_pct, total
        ("INV-1001", "2026-04-10", "CP-001", "Claims Processing",  12000,  5.50,  5, 8,  None),  # FAIL: price + discount
        ("INV-1001", "2026-04-10", "AP-002", "Appeals Processing",   200, 12.00,  0, 8,  None),  # PASS
        ("INV-1001", "2026-04-10", "EL-003", "Eligibility Checks",  8500,  3.25,  8, 8,  None),  # FAIL: discount
        ("INV-1001", "2026-04-10", "DM-004", "Data Migration",       150, 45.00,  0, 8,  None),  # FAIL: discount
    ]

    def calc_total(qty, price, disc_pct, tax_pct):
        return round(qty * price * (1 - disc_pct / 100) * (1 + tax_pct / 100), 2)

    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["invoice_id", "invoice_date", "sku", "description",
                         "quantity", "unit_price", "discount_percent", "tax_percent", "total_amount"])
        for r in rows:
            inv_id, date, sku, desc, qty, price, disc, tax, _ = r
            total = calc_total(qty, price, disc, tax)
            writer.writerow([inv_id, date, sku, desc, qty, price, disc, tax, total])

    print(f"  Created: {path}")
    print()
    print("  Expected (authoritative) values:")
    for sku, t in FINAL_TERMS.items():
        print(f"    {sku}: unit_price=${t['unit_price']:.2f}  discount={t['discount']}%  tax={t['tax']}%")


# ── Main ────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating synthetic dataset...")
    make_pdf()
    make_excel()
    make_docx()
    make_email()
    make_invoice()
    print("\nDone. Files written to data/contracts/ and data/invoices/")
