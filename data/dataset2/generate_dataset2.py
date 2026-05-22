"""
Dataset 2 — NovaTech Solutions SaaS vendor contract audit
Scenario: Cloud services (storage, API, support, licenses)
Intentional invoice mismatches to test the auditor.

Ground truth expected values (after amendments effective 2026-01-01):
  SRV-101  Storage        unit_price=0.06  discount=5%   tax=10%
  API-202  API Calls      unit_price=2.50  discount=5%   tax=10%   ← email addendum adds 5%
  SUP-303  Premium Support unit_price=150  discount=10%  tax=10%
  LIC-404  Software Lic   unit_price=45    discount=20%  tax=10%   ← docx amendment bumps 15→20%

Invoice mismatches (what the vendor actually billed):
  SRV-101  unit_price=0.08 (old price, ignores amendment) → FAIL
  API-202  discount=0%    (misses email addendum 5%)      → FAIL
  SUP-303  all correct                                    → PASS
  LIC-404  discount=15%   (old rate, ignores amendment)   → FAIL
"""
import os
import csv
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

OUT = os.path.dirname(__file__)
CONTRACTS = os.path.join(OUT, "contracts")
INVOICES  = os.path.join(OUT, "invoices")

# ── 1. Invoice CSV ────────────────────────────────────────────────────────────
def make_invoice():
    rows = [
        # invoice_id, date,        sku,     description,        qty,    unit_price, disc%,  tax%, total
        ["INV-2001", "2026-04-15", "SRV-101", "Cloud Storage (GB/mo)",  500000, 0.08, 5,   10,  0],
        ["INV-2001", "2026-04-15", "API-202", "API Calls (per 1k)",     80000,  2.50, 0,   10,  0],
        ["INV-2001", "2026-04-15", "SUP-303", "Premium Support (hrs)",  120,    150,  10,  10,  0],
        ["INV-2001", "2026-04-15", "LIC-404", "Software Licenses",      300,    45,   15,  10,  0],
    ]
    # compute totals: qty * price * (1 - disc/100) * (1 + tax/100)
    for r in rows:
        qty, price, disc, tax = r[5], r[6], r[7], r[8]  # wait indices are wrong
    # redo with correct indices
    header = ["invoice_id","invoice_date","sku","description","quantity","unit_price","discount_percent","tax_percent","total_amount"]
    data = []
    for r in rows:
        inv_id, date, sku, desc, qty, price, disc, tax, _ = r
        total = round(qty * price * (1 - disc/100) * (1 + tax/100), 2)
        data.append([inv_id, date, sku, desc, qty, price, disc, tax, total])

    path = os.path.join(INVOICES, "invoice_INV-2001.csv")
    with open(path, "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(header)
        w.writerows(data)
    print(f"  [invoice] {path}")
    for d in data:
        print(f"    {d[2]:8s}  unit={d[5]}  disc={d[6]}%  total={d[8]}")

# ── 2. Pricing XLSX ───────────────────────────────────────────────────────────
def make_xlsx():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Base Pricing ──
    ws1 = wb.active
    ws1.title = "Base Pricing"
    hdr_fill  = PatternFill("solid", fgColor="1F3864")
    hdr_font  = Font(bold=True, color="FFFFFF", size=11)
    bold      = Font(bold=True, size=10)
    thin      = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    ws1["A1"] = "NOVATECH SOLUTIONS — BASE PRICING SCHEDULE"
    ws1["A1"].font = Font(bold=True, size=13, color="1F3864")
    ws1.merge_cells("A1:G1")
    ws1["A2"] = "Effective: 2025-01-01 | Agreement: SaaS-MSA-2025-NVT-0012 | Currency: USD"
    ws1["A2"].font = Font(italic=True, size=9, color="595959")
    ws1.merge_cells("A2:G2")
    ws1.append([])

    cols = ["SKU","Service Description","Category","Unit Price","Std Discount (%)","Tax Rate (%)","Notes"]
    ws1.append(cols)
    for i, c in enumerate(cols, 1):
        cell = ws1.cell(row=4, column=i)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin
        cell.alignment = Alignment(horizontal="center")

    base_rows = [
        ("SRV-101","Cloud Storage (GB/mo)","Infrastructure",  0.08, 5,  10, "Per GB per month"),
        ("API-202","API Calls (per 1k)",   "Platform",        2.50, 0,  10, "Per 1,000 API calls"),
        ("SUP-303","Premium Support (hrs)","Support",         150,  10, 10, "Dedicated SLA support"),
        ("LIC-404","Software Licenses",    "Licensing",       45,   15, 10, "Per-seat monthly license"),
    ]
    alt_fill = PatternFill("solid", fgColor="EBF3FB")
    for i, row in enumerate(base_rows):
        ws1.append(list(row))
        r = ws1.max_row
        for col in range(1, 8):
            cell = ws1.cell(row=r, column=col)
            if i % 2 == 1:
                cell.fill = alt_fill
            cell.border = thin

    for letter in "ABCDEFG":
        ws1.column_dimensions[letter].width = 22

    # ── Sheet 2: Discounts & Amendments ──
    ws2 = wb.create_sheet("Discounts & Amendments")
    ws2["A1"] = "DISCOUNT SCHEDULE & AMENDMENT LOG"
    ws2["A1"].font = Font(bold=True, size=13, color="1F3864")
    ws2.merge_cells("A1:G1")
    ws2["A2"] = "This sheet records all active discounts and amendment history. Amendments supersede base contract terms."
    ws2["A2"].font = Font(italic=True, size=9, color="595959")
    ws2.merge_cells("A2:G2")
    ws2.append([])

    disc_cols = ["SKU","Service","Discount Type","Discount (%)","Effective Date","Expiry Date","Source Document"]
    ws2.append(disc_cols)
    for i, c in enumerate(disc_cols, 1):
        cell = ws2.cell(row=4, column=i)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin
        cell.alignment = Alignment(horizontal="center")

    disc_rows = [
        ("SRV-101","Cloud Storage",   "Standard",          5,  "2025-01-01","Ongoing",       "master_saas_contract.pdf"),
        ("SRV-101","Cloud Storage",   "Amendment Override", 5,  "2026-01-01","Ongoing",       "saas_amendment_q1.docx"),
        ("API-202","API Calls",       "Standard",          0,  "2025-01-01","2025-12-31",    "master_saas_contract.pdf"),
        ("API-202","API Calls",       "Email Addendum",    5,  "2026-01-01","Ongoing",       "novatech_email_addendum.eml"),
        ("SUP-303","Premium Support", "Standard",          10, "2025-01-01","Ongoing",       "master_saas_contract.pdf"),
        ("LIC-404","Software Lic",    "Standard",          15, "2025-01-01","2025-12-31",    "master_saas_contract.pdf"),
        ("LIC-404","Software Lic",    "Amendment Override",20, "2026-01-01","Ongoing",       "saas_amendment_q1.docx"),
    ]
    amend_fill = PatternFill("solid", fgColor="FFF2CC")
    for i, row in enumerate(disc_rows):
        ws2.append(list(row))
        r = ws2.max_row
        for col in range(1, 8):
            cell = ws2.cell(row=r, column=col)
            if "Amendment" in str(row[2]) or "Addendum" in str(row[2]):
                cell.fill = amend_fill
            elif i % 2 == 1:
                cell.fill = alt_fill
            cell.border = thin

    for letter in "ABCDEFG":
        ws2.column_dimensions[letter].width = 22

    # ── Sheet 3: Volume Tiers ──
    ws3 = wb.create_sheet("Volume Tiers")
    ws3["A1"] = "VOLUME DISCOUNT TIERS"
    ws3["A1"].font = Font(bold=True, size=13, color="1F3864")
    ws3.merge_cells("A1:F1")
    ws3.append([])
    tier_cols = ["SKU","Service","< 10k units","10k – 99k units","100k – 499k units","≥ 500k units"]
    ws3.append(tier_cols)
    for i, c in enumerate(tier_cols, 1):
        cell = ws3.cell(row=3, column=i)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin
    tier_rows = [
        ("SRV-101","Cloud Storage",  0, 0, 1, 2),
        ("API-202","API Calls",      0, 1, 2, 3),
        ("SUP-303","Premium Support",0, 0, 0, 1),
        ("LIC-404","Software Lic",   0, 1, 2, 3),
    ]
    for row in tier_rows:
        ws3.append(list(row))
    for letter in "ABCDEF":
        ws3.column_dimensions[letter].width = 20

    path = os.path.join(CONTRACTS, "saas_pricing.xlsx")
    wb.save(path)
    print(f"  [xlsx]    {path}")

# ── 3. DOCX Amendment ─────────────────────────────────────────────────────────
def make_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("AMENDMENT NO. 1 TO SAAS MASTER SERVICES AGREEMENT", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph("Agreement Reference: SaaS-MSA-2025-NVT-0012")
    doc.add_paragraph("Amendment Effective Date: January 1, 2026")
    doc.add_paragraph(
        'This Amendment No. 1 is entered into by Acme Corp (Client) '
        'and NovaTech Solutions (Provider) and amends the SaaS Master Services Agreement '
        'dated January 1, 2025. Except as modified herein, all other terms remain in effect.'
    )

    doc.add_heading("Section 1 — Purpose", level=1)
    doc.add_paragraph(
        "Following a competitive pricing review and updated infrastructure costs, the parties "
        "agree to revise unit pricing and discount terms for select SKUs effective Q1 2026."
    )

    doc.add_heading("Section 2 — Revised Unit Pricing", level=1)
    doc.add_paragraph(
        "Effective January 1, 2026, the unit price for SKU SRV-101 (Cloud Storage) is REVISED "
        "from $0.08 per GB to $0.06 per GB per month. This amendment supersedes the base contract "
        "price stated in the Master Services Agreement and in saas_pricing.xlsx (Base Pricing sheet). "
        "All other SKU unit prices remain unchanged."
    )

    doc.add_heading("Section 3 — Revised Discount Schedule", level=1)
    doc.add_paragraph(
        "Effective January 1, 2026, the following discount changes apply:"
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["SKU","Service","Old Discount","New Discount (effective 2026-01-01)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    changes = [
        ("LIC-404","Software Licenses","15%","20%"),
    ]
    for row in changes:
        r = tbl.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_paragraph(
        "The revised discounts supersede the rates in saas_pricing.xlsx (Base Pricing and "
        "Discounts & Amendments sheets). Invoices submitted after January 1, 2026 must reflect "
        "the updated rates."
    )

    doc.add_heading("Section 4 — No Other Changes", level=1)
    doc.add_paragraph(
        "Except as expressly modified above, all terms and conditions of SaaS-MSA-2025-NVT-0012, "
        "including Exhibit A (saas_pricing.xlsx), remain in full force and effect."
    )

    doc.add_heading("Section 5 — Signatures", level=1)
    doc.add_paragraph("Acme Corp:            ________________________    Date: 2026-01-01")
    doc.add_paragraph("NovaTech Solutions:   ________________________    Date: 2026-01-01")

    path = os.path.join(CONTRACTS, "saas_amendment_q1.docx")
    doc.save(path)
    print(f"  [docx]    {path}")

# ── 4. Email Addendum ─────────────────────────────────────────────────────────
def make_eml():
    body = """\
From: Marcus Lee <marcus.lee@novatech.io>
To: procurement@acmecorp.com; legal@acmecorp.com
CC: contracts@novatech.io
Subject: SaaS-MSA-2025-NVT-0012 — API-202 Discount Addendum (Q1 2026)
Date: Mon, 15 Dec 2025 09:30:00 -0500

Hi Team,

Following our Q4 review call, we are pleased to confirm a 5% discount on SKU API-202
(API Calls) effective January 1, 2026.

Summary of agreed terms:

  SKU:              API-202 — API Calls (per 1,000 calls)
  Discount Change:  0% (current) → 5% (effective 2026-01-01)
  Unit Price:       $2.50 per 1,000 calls (unchanged)
  Tax Rate:         10% (unchanged)

This email serves as a binding addendum to SaaS Master Services Agreement
SaaS-MSA-2025-NVT-0012 and supersedes any prior discount schedule for API-202.
Please reference this email as "Email Addendum — API-202 Discount, dated 2025-12-15"
in all future invoices and audits.

Please reply to confirm acceptance.

Best regards,
Marcus Lee
Director of Partnerships — NovaTech Solutions
Ph: +1 (512) 334-8821

---
This email and any attachments are confidential.
"""
    path = os.path.join(CONTRACTS, "novatech_email_addendum.eml")
    with open(path, "w") as f:
        f.write(body)
    print(f"  [eml]     {path}")

# ── 5. PDF Base Contract ──────────────────────────────────────────────────────
def make_pdf():
    path = os.path.join(CONTRACTS, "master_saas_contract.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    h1  = styles["Heading1"]
    h2  = styles["Heading2"]
    body= styles["BodyText"]
    body.spaceAfter = 8

    title_style = ParagraphStyle("title", parent=styles["Title"],
                                 fontSize=16, textColor=colors.HexColor("#1F3864"),
                                 spaceAfter=12)
    sub_style   = ParagraphStyle("sub", parent=body,
                                 fontSize=9, textColor=colors.grey, spaceAfter=16)

    story = []
    story.append(Paragraph("SAAS MASTER SERVICES AGREEMENT", title_style))
    story.append(Paragraph(
        "Agreement ID: SaaS-MSA-2025-NVT-0012 | Effective: January 1, 2025 | "
        "Parties: Acme Corp (Client) and NovaTech Solutions (Provider)",
        sub_style))

    story.append(Paragraph("1. SCOPE OF SERVICES", h2))
    story.append(Paragraph(
        "Provider shall deliver cloud infrastructure, API platform, technical support, and "
        "software licensing services to Client as described in Exhibit A (Pricing Schedule). "
        "Services are subject to availability, usage limits, and SLA commitments defined herein.",
        body))

    story.append(Paragraph("2. PRICING AND FEES", h2))
    story.append(Paragraph(
        "Fees are calculated monthly based on actual usage. Unit prices, standard discounts, "
        "and tax rates are defined in Exhibit A. Client shall be invoiced within 5 business days "
        "of each calendar month end.",
        body))

    tbl_data = [
        ["SKU", "Service", "Unit Price", "Std Discount", "Tax Rate"],
        ["SRV-101", "Cloud Storage (GB/mo)",   "$0.08/GB",    "5%",  "10%"],
        ["API-202", "API Calls (per 1k)",       "$2.50/1k",   "0%",  "10%"],
        ["SUP-303", "Premium Support (hrs)",    "$150.00/hr", "10%", "10%"],
        ["LIC-404", "Software Licenses",        "$45.00/seat","15%", "10%"],
    ]
    tbl = Table(tbl_data, colWidths=[0.9*inch, 2.2*inch, 1.1*inch, 1.1*inch, 0.9*inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0), colors.HexColor("#1F3864")),
        ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
        ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EBF3FB")]),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ("ALIGN",       (2,0), (-1,-1), "CENTER"),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",  (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))

    story.append(Paragraph("3. AMENDMENTS", h2))
    story.append(Paragraph(
        "This Agreement may be amended by mutual written consent. Amendments shall be numbered "
        "sequentially (Amendment No. 1, No. 2, etc.) and become effective on their stated date. "
        "Email communications acknowledged by both parties constitute binding addenda when "
        "explicitly referencing this Agreement ID.",
        body))

    story.append(Paragraph("4. TAXES", h2))
    story.append(Paragraph(
        "All fees are subject to a 10% tax rate unless otherwise specified. Client is responsible "
        "for applicable state and local taxes beyond the contractual rate.",
        body))

    story.append(Paragraph("5. TERM AND TERMINATION", h2))
    story.append(Paragraph(
        "This Agreement commences January 1, 2025 and continues for an initial term of 24 months. "
        "Either party may terminate with 60 days written notice after the initial term.",
        body))

    story.append(Paragraph("6. DISPUTE RESOLUTION", h2))
    story.append(Paragraph(
        "Disputes shall first be escalated to executive leadership of both parties. If unresolved "
        "within 30 days, disputes shall be submitted to binding arbitration under AAA Commercial Rules.",
        body))

    story.append(Spacer(1, 24))
    story.append(Paragraph("SIGNATURES", h2))
    sig_data = [
        ["Acme Corp",              "NovaTech Solutions"],
        ["________________________","________________________"],
        ["Authorized Signatory",   "Authorized Signatory"],
        ["Date: 2025-01-01",       "Date: 2025-01-01"],
    ]
    sig_tbl = Table(sig_data, colWidths=[3*inch, 3*inch])
    sig_tbl.setStyle(TableStyle([
        ("FONTNAME",  (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0), (-1,-1), 10),
        ("ALIGN",     (0,0), (-1,-1), "LEFT"),
        ("TOPPADDING",(0,0), (-1,-1), 4),
    ]))
    story.append(sig_tbl)

    doc.build(story)
    print(f"  [pdf]     {path}")

# ── Run all ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating Dataset 2 — NovaTech SaaS Vendor...")
    make_invoice()
    make_xlsx()
    make_docx()
    make_eml()
    make_pdf()
    print("\nDone! Expected audit results:")
    print("  SRV-101  unit_price  0.06 vs 0.08  FAIL  (amendment cut price)")
    print("  SRV-101  discount    5%   vs 5%    PASS")
    print("  SRV-101  tax         10%  vs 10%   PASS")
    print("  SRV-101  total       expected=25350 vs billed=38000  FAIL")
    print("  API-202  unit_price  2.50 vs 2.50  PASS")
    print("  API-202  discount    5%   vs 0%    FAIL  (email addendum missed)")
    print("  API-202  tax         10%  vs 10%   PASS")
    print("  API-202  total       expected=209000 vs billed=220000  FAIL")
    print("  SUP-303  all fields  correct       PASS x4")
    print("  LIC-404  unit_price  45   vs 45    PASS")
    print("  LIC-404  discount    20%  vs 15%   FAIL  (amendment increase missed)")
    print("  LIC-404  tax         10%  vs 10%   PASS")
    print("  LIC-404  total       expected=11880 vs billed=12881.25  FAIL")
